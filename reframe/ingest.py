#!/usr/bin/env python3
"""Load a resume from heterogeneous sources into one clean canonical string.

The returned string is the single source of truth for everything downstream:
segmentation offsets index it, and the grounding guard normalizes against it, so
it must be produced ONCE, here, before anything else looks at the resume.

PII discipline: parsing happens in memory only. Nothing is written to disk or
logs, and errors never carry resume content - an :class:`IngestError` is a
one-line remediation hint, not a stack trace over the bytes.

PDF/DOCX support is lazy and optional: the libraries are imported inside the
loader functions so the package (and every test) imports with zero extra deps.
"""
import io
import os
import sys

# Characters that would desync offset math or the guard's normalization.
_ZERO_WIDTH = ("﻿", "​", "⁠")  # BOM / ZWSP / word-joiner
_NBSP = (" ", " ")  # non-breaking / narrow non-breaking space


class IngestError(Exception):
    """A resume could not be loaded. Message is a remediation hint, never text."""

    def __init__(self, fmt, hint):
        super().__init__(f"cannot read {fmt} resume: {hint}")
        self.fmt = fmt
        self.hint = hint


def clean_text(raw):
    """Normalize a raw resume string into the canonical form used everywhere.

    Newlines -> LF, drop zero-width chars, fold non-breaking spaces to a normal
    space, and right-strip each line WITHOUT changing the line count (so offsets
    and the guard's normalization index the same clean string).
    """
    if raw is None:
        return ""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    for ch in _ZERO_WIDTH:
        text = text.replace(ch, "")
    for ch in _NBSP:
        text = text.replace(ch, " ")
    return "\n".join(line.rstrip() for line in text.split("\n"))


def load_text_stream(fh):
    """Read an already-open text handle (file, StringIO) and clean it."""
    return clean_text(fh.read())


def _load_pdf(path):
    try:
        import pypdf
    except ImportError:
        raise IngestError("pdf", "pip install pypdf") from None
    try:
        reader = pypdf.PdfReader(path)
        pages = [(pg.extract_text() or "") for pg in reader.pages]
    except Exception:  # never surface a parser trace that may echo bytes
        raise IngestError("pdf", "the file is not a readable PDF") from None
    return clean_text("\n\n".join(pages))


def _load_docx(path):
    try:
        import docx
    except ImportError:
        raise IngestError("docx", "pip install python-docx") from None
    try:
        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
    except Exception:
        raise IngestError("docx", "the file is not a readable .docx") from None
    return clean_text("\n".join(parts))


def load_resume(source=None, kind=None):
    """Load a resume from a path, an open handle, or stdin, into clean text.

    ``source`` may be None or '-' (read stdin), a filesystem path, or an open
    text handle. ``kind`` ('text'/'pdf'/'docx') overrides suffix detection.
    """
    if source is None or source == "-":
        return clean_text(sys.stdin.read())
    if isinstance(source, io.IOBase) or hasattr(source, "read"):
        return load_text_stream(source)

    path = str(source)
    fmt = kind or os.path.splitext(path)[1].lstrip(".").lower()
    if fmt == "pdf":
        return _load_pdf(path)
    if fmt == "docx":
        return _load_docx(path)
    # text / md / anything else: read as UTF-8 text, tolerating bad bytes.
    with open(path, encoding="utf-8", errors="replace") as fh:
        return clean_text(fh.read())
